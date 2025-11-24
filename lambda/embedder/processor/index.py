"""
Lambda 1: Nova MME Processor

Triggered by S3 upload via Step Functions.
- Extracts metadata from S3 object
- Formats request for Nova MME async invocation at 3072 dimensions
- Starts async job
- Returns invocation ARN and metadata for Step Functions
"""

import json
import os
import boto3
from datetime import datetime
from typing import Dict, Any, List
from urllib.parse import unquote_plus
import tempfile

# PyMuPDF import (optional for testing)
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("Warning: PyMuPDF not installed, PDF support disabled")

# python-docx import (optional for testing)
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not installed, .docx support disabled")

# Initialize clients
s3_client = boto3.client('s3')
bedrock_runtime = boto3.client('bedrock-runtime')

# Environment variables
EMBEDDING_DIMENSION = int(os.environ.get('EMBEDDING_DIMENSION', '3072'))
MODEL_ID = os.environ.get('MODEL_ID', 'amazon.nova-2-multimodal-embeddings-v1:0')
OUTPUT_BUCKET = os.environ['OUTPUT_BUCKET']

# File type mappings
IMAGE_FORMATS = {
    '.png': 'png', '.jpg': 'jpeg', '.jpeg': 'jpeg',
    '.gif': 'gif', '.webp': 'webp'
}
VIDEO_FORMATS = {
    '.mp4': 'mp4', '.mov': 'mov', '.mkv': 'mkv', '.webm': 'webm',
    '.flv': 'flv', '.mpeg': 'mpeg', '.mpg': 'mpg', '.wmv': 'wmv', '.3gp': '3gp'
}
AUDIO_FORMATS = {
    '.mp3': 'mp3', '.wav': 'wav', '.ogg': 'ogg'
}
TEXT_FORMATS = ['.txt', '.md', '.json', '.csv']
DOCUMENT_FORMATS = ['.docx', '.doc']  # Extracted as text
# PDFs are handled separately - converted to images then processed with DOCUMENT_IMAGE
# Google Docs format (.gdoc) is a pointer file, not the actual document - users must export to .docx first


def handler(event, context):
    """
    Main handler for Nova MME Processor Lambda
    
    Args:
        event: Contains 'bucket' and 'key' from Step Functions
        context: Lambda context
    
    Returns:
        Dict with invocationArn and metadata for next Lambda
    """
    try:
        # Extract S3 info from event
        bucket = event['bucket']
        key = event['key']
        
        # URL-decode the key if needed (defense in depth)
        key = unquote_plus(key)
        
        print(f"Processing file: s3://{bucket}/{key}")
        
        # Extract metadata from S3 object
        metadata = extract_s3_metadata(bucket, key)
        print(f"Extracted metadata: {json.dumps(metadata)}")
        
        # Determine file type
        file_extension = os.path.splitext(key)[1].lower()
        
        # Special handling for Word documents - extract text and process as TEXT
        if file_extension in DOCUMENT_FORMATS:
            if not DOCX_SUPPORT:
                raise Exception("python-docx not installed, cannot process .docx files")
            
            print(f"Extracting text from Word document...")
            
            # Download document to temp file
            local_path = f"/tmp/{os.path.basename(key)}"
            s3_client.download_file(bucket, key, local_path)
            
            # Extract text
            text_content = extract_docx_text(local_path)
            print(f"Extracted {len(text_content)} characters from document")
            
            # Upload extracted text to temp location
            text_key = f"docx-text/{metadata['objectId']}.txt"
            s3_client.put_object(
                Bucket=bucket,
                Key=text_key,
                Body=text_content.encode('utf-8'),
                ContentType='text/plain'
            )
            print(f"Uploaded extracted text to s3://{bucket}/{text_key}")
            
            # Update metadata to point to text file
            metadata['originalDocUri'] = metadata['sourceS3Uri']
            metadata['sourceS3Uri'] = f"s3://{bucket}/{text_key}"
            metadata['isDocx'] = True
            metadata['originalFileName'] = metadata['fileName']
            
            # Process as text
            model_input = create_model_input(bucket, text_key, '.txt')
            print(f"Created model input for extracted text")
            
            # Start async invocation
            output_s3_uri = f"s3://{OUTPUT_BUCKET}/{metadata['objectId']}/"
            invocation_arn = start_async_invocation(model_input, output_s3_uri)
            print(f"Started async invocation: {invocation_arn}")
            
            return {
                'statusCode': 200,
                'invocationArn': invocation_arn,
                'outputS3Uri': output_s3_uri,
                'metadata': metadata,
                'status': 'IN_PROGRESS'
            }
        
        # Special handling for PDFs - convert to images and process each page
        elif file_extension == '.pdf':
            print(f"Converting PDF to images...")
            image_uris = convert_pdf_to_images(bucket, key, metadata['objectId'])
            print(f"Converted PDF to {len(image_uris)} images")
            
            # Process ALL pages - return list of jobs for Step Functions to handle
            pdf_pages = []
            
            for page_num, image_uri in enumerate(image_uris, start=1):
                # Create model input for this page
                image_key = image_uri.replace(f"s3://{bucket}/", "")
                model_input = create_model_input(bucket, image_key, '.png')
                model_input['segmentedEmbeddingParams']['image']['detailLevel'] = 'DOCUMENT_IMAGE'
                
                # Create page-specific metadata
                page_metadata = metadata.copy()
                page_metadata['isPdf'] = True
                page_metadata['totalPages'] = len(image_uris)
                page_metadata['processedPage'] = page_num
                page_metadata['objectId'] = f"{metadata['objectId']}_page_{page_num}"
                
                # Start async invocation for this page
                output_s3_uri = f"s3://{OUTPUT_BUCKET}/{page_metadata['objectId']}/"
                invocation_arn = start_async_invocation(model_input, output_s3_uri)
                print(f"Started async invocation for page {page_num}: {invocation_arn}")
                
                pdf_pages.append({
                    'invocationArn': invocation_arn,
                    'outputS3Uri': output_s3_uri,
                    'metadata': page_metadata
                })
            
            # Return all pages - Step Functions will need to process each
            # For now, return first page in standard format for compatibility
            return {
                'statusCode': 200,
                'invocationArn': pdf_pages[0]['invocationArn'],
                'outputS3Uri': pdf_pages[0]['outputS3Uri'],
                'metadata': pdf_pages[0]['metadata'],
                'status': 'IN_PROGRESS',
                'pdfPages': pdf_pages,  # All pages for potential batch processing
                'totalPages': len(image_uris)
            }
        else:
            # Regular file processing
            model_input = create_model_input(bucket, key, file_extension)
            print(f"Created model input for type: {metadata['fileType']}")
            
            # Start async invocation
            # Output URI must point to a directory (end with /)
            output_s3_uri = f"s3://{OUTPUT_BUCKET}/{metadata['objectId']}/"
            invocation_arn = start_async_invocation(model_input, output_s3_uri)
            print(f"Started async invocation: {invocation_arn}")
            
            # Return data for Step Functions
            return {
                'statusCode': 200,
                'invocationArn': invocation_arn,
                'outputS3Uri': output_s3_uri,
                'metadata': metadata,
                'status': 'IN_PROGRESS'
            }
        
    except Exception as e:
        print(f"Error processing file: {str(e)}")
        return {
            'statusCode': 500,
            'error': str(e),
            'status': 'FAILED'
        }


def extract_s3_metadata(bucket: str, key: str) -> Dict[str, Any]:
    """
    Extract metadata from S3 object
    
    Returns:
        Dict with sourceS3Uri, fileName, fileType, fileSize, uploadTimestamp, etc.
    """
    # Get object metadata
    response = s3_client.head_object(Bucket=bucket, Key=key)
    
    # Generate unique object ID
    object_id = key.replace('/', '_').replace('.', '_') + '_' + datetime.now().strftime('%Y%m%d%H%M%S')
    
    metadata = {
        'sourceS3Uri': f"s3://{bucket}/{key}",
        'fileName': os.path.basename(key),
        'fileType': os.path.splitext(key)[1].lower(),
        'fileSize': response['ContentLength'],
        'uploadTimestamp': response['LastModified'].isoformat(),
        'contentType': response.get('ContentType', 'unknown'),
        'objectId': object_id
    }
    
    return metadata


def create_model_input(bucket: str, key: str, file_extension: str) -> Dict[str, Any]:
    """
    Create model input JSON based on file type
    
    Returns:
        Dict with properly formatted Nova MME async request
    """
    s3_uri = f"s3://{bucket}/{key}"
    
    # Base structure
    model_input = {
        "schemaVersion": "nova-multimodal-embed-v1",
        "taskType": "SEGMENTED_EMBEDDING",
        "segmentedEmbeddingParams": {
            "embeddingPurpose": "GENERIC_INDEX",
            "embeddingDimension": EMBEDDING_DIMENSION
        }
    }
    
    # Add modality-specific configuration
    if file_extension in IMAGE_FORMATS:
        model_input["segmentedEmbeddingParams"]["image"] = {
            "format": IMAGE_FORMATS[file_extension],
            "source": {
                "s3Location": {"uri": s3_uri}
            },
            # Use DOCUMENT_IMAGE for better text/diagram interpretation
            "detailLevel": "DOCUMENT_IMAGE"
        }
    
    elif file_extension in VIDEO_FORMATS:
        model_input["segmentedEmbeddingParams"]["video"] = {
            "format": VIDEO_FORMATS[file_extension],
            "source": {
                "s3Location": {"uri": s3_uri}
            },
            "embeddingMode": "AUDIO_VIDEO_COMBINED",
            "segmentationConfig": {
                "durationSeconds": 5
            }
        }
    
    elif file_extension in AUDIO_FORMATS:
        model_input["segmentedEmbeddingParams"]["audio"] = {
            "format": AUDIO_FORMATS[file_extension],
            "source": {
                "s3Location": {"uri": s3_uri}
            },
            "segmentationConfig": {
                "durationSeconds": 5
            }
        }
    
    elif file_extension in TEXT_FORMATS:
        model_input["segmentedEmbeddingParams"]["text"] = {
            "truncationMode": "END",
            "source": {
                "s3Location": {"uri": s3_uri}
            },
            "segmentationConfig": {
                "maxLengthChars": 32000
            }
        }
    
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
    
    return model_input


def extract_docx_text(docx_path: str) -> str:
    """
    Extract plain text from .docx file
    
    Args:
        docx_path: Local path to .docx file
    
    Returns:
        Extracted text content
    
    Note: .doc (old Word format) requires conversion to .docx first.
    For demo purposes, we only support .docx natively.
    """
    try:
        doc = Document(docx_path)
        
        # Extract text from all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
        
        # Combine all text
        all_text = '\n\n'.join(paragraphs)
        if table_text:
            all_text += '\n\n--- Tables ---\n\n' + '\n'.join(table_text)
        
        return all_text
        
    except Exception as e:
        print(f"Error extracting text from .docx: {e}")
        raise


def convert_pdf_to_images(bucket: str, key: str, object_id: str) -> List[str]:
    """
    Convert PDF pages to images and upload to S3
    
    Returns:
        List of S3 URIs for the converted images
    """
    if not PDF_SUPPORT:
        raise RuntimeError("PDF support not available - PyMuPDF not installed")
    
    # Download PDF from S3
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
        s3_client.download_fileobj(bucket, key, tmp_pdf)
        pdf_path = tmp_pdf.name
    
    try:
        # Open PDF with PyMuPDF
        pdf_document = fitz.open(pdf_path)
        image_uris = []
        
        # Convert each page to image
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Render page to image at high resolution (300 DPI for DOCUMENT_IMAGE)
            pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
            
            # Convert to PNG bytes
            img_bytes = pix.tobytes("png")
            
            # Upload to S3 in a permanent location (needed for chatbot to fetch later)
            image_key = f"pdf-pages/{object_id}/page_{page_num + 1}.png"
            s3_client.put_object(
                Bucket=bucket,
                Key=image_key,
                Body=img_bytes,
                ContentType='image/png'
            )
            
            image_uris.append(f"s3://{bucket}/{image_key}")
            print(f"Converted PDF page {page_num + 1} to {image_key}")
        
        pdf_document.close()
        return image_uris
        
    finally:
        # Clean up temp file
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)


def start_async_invocation(model_input: Dict[str, Any], output_s3_uri: str) -> str:
    """
    Start async invocation with Bedrock
    
    Returns:
        Invocation ARN for status checking
    """
    response = bedrock_runtime.start_async_invoke(
        modelId=MODEL_ID,
        modelInput=model_input,
        outputDataConfig={
            "s3OutputDataConfig": {
                "s3Uri": output_s3_uri
            }
        }
    )
    
    return response['invocationArn']
